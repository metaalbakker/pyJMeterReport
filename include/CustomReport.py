import mammoth
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime
import os
import re
import collections
from operator import itemgetter

pagecount_number=3

def create_html_from_docx(docx_filename):
    import mammoth
    with open(docx_filename, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value  # The generated HTML
        # messages = result.messages                     # Any messages, such as warnings during conversion
    docx_file.close()

    html_filename = docx_filename.replace('.docx', '.html')  # HTML filename  based on docx_filename

    with open(html_filename, "w") as html_file:
        html_file.write(html)

    html_file.close()

    return html_filename

def make_report(data, filename, stats_json, images_list, results_dir):


    home_path = os.getcwd()
    work_path = str(os.path.dirname(os.path.abspath(__file__)))

    # print(work_path)

    document = Document(work_path + '/docx_template/pt_testreport_template.docx')
    document.add_picture(work_path + '/docx_template/picture.png', width=Cm(9.0))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style = document.styles['Heading 1']
    font = style.font
    font.size = Pt(12)

    ################################
    ### Section: Vrijgave Advies ###
    ################################
    document.add_heading("Vrijgave Advies", level=1)
    par_va = document.add_paragraph()
    par_va.add_run(
        "Bij deze geeft team LSP een positief/negatief (verwijder wat niet van toepassing is) vrijgave advies voor de in-productiename van").font.size = Pt(
        10)

    #######################################################
    ## Section: Performance Eisen & Gebruiksverwachting ###
    #######################################################
    document.add_heading("Doel van de test", level=1)
    par_pe = document.add_paragraph()
    par_pe.add_run(
        "<Beschrijf hier het doel van de test en de performance eisen en/of gebruiksverwachting>").font.size = Pt(10)

    #######################################################
    ## Section: Testuitvoering                          ###
    #######################################################
    document.add_heading("Testuitvoering", level=1)
    par_pe = document.add_paragraph()
    par_pe.add_run("<Beschrijf hier details over het testobject en hoe de test is uitgevoerd>").font.size = Pt(10)

    ################################
    ### Section: Test Resultaten ###
    ################################
    # document.add_page_break()
    document.add_heading("Test Resultaten", level=1)

    metrics_= set()
    statuses_ = set()

    for metingid_,status_dict_ in stats_json[data].items():
        for status_,metric_dict_ in status_dict_.items():
            statuses_.add(status_)
            for metric_,statistics_dict_ in metric_dict_.items():
                metrics_.add(metric_)

    for unique_metric_ in metrics_:


        tbl_heading = document.add_heading(str(unique_metric_), level=2)
        table_res = document.add_table(rows=1, cols=9)
        table_res.style = 'Light List Accent 1'
        table_res.alignment = WD_TABLE_ALIGNMENT.LEFT

        ## Write table headers
        hdr_cells = table_res.rows[0].cells
        ### Naam van dekolommen


        hdr_cells[0].text = "Meting ID"
        hdr_cells[1].text = "Status"
        hdr_cells[2].text = "Volume"
        hdr_cells[3].text = "Avg"
        hdr_cells[4].text = "Min"
        hdr_cells[5].text = "Max"
        hdr_cells[6].text = "Stddev"
        hdr_cells[7].text = "Perc95"

        for unique_status_ in statuses_:

            for metingid_,status_dict_ in stats_json[data].items():
                if unique_status_ in status_dict_:

                    statistics_dict_ = status_dict_[unique_status_][unique_metric_]
                    # Fill table
                    row_cells = table_res.add_row().cells  # New row
                    row_cells[0].text = re.sub("(.{50})", "\\1\n", str(metingid_), 0, re.DOTALL)  # Insert a newline character every 64 characters
                    row_cells[1].text = re.sub("(.{50})", "\\1\n", str(unique_status_), 0, re.DOTALL)
                    row_cells[2].text = re.sub("(.{50})", "\\1\n", str(statistics_dict_["Count"]), 0, re.DOTALL)
                    row_cells[3].text = re.sub("(.{50})", "\\1\n", str(statistics_dict_["Mean"]), 0, re.DOTALL)
                    row_cells[4].text = re.sub("(.{50})", "\\1\n", str(statistics_dict_["Min"]), 0, re.DOTALL)
                    row_cells[5].text = re.sub("(.{50})", "\\1\n", str(statistics_dict_["Max"]), 0,re.DOTALL)
                    row_cells[6].text = re.sub("(.{50})", "\\1\n", str(statistics_dict_["StDev"]), 0, re.DOTALL)
                    row_cells[7].text = re.sub("(.{50})", "\\1\n", str(statistics_dict_["Percentile.95"]), 0,re.DOTALL)

    ##########################
    ### Section: Conclusie ###
    ##########################
    document.add_heading("Conclusie", level=1)
    par_con = document.add_paragraph()
    par_con.add_run(
        "<Beschrijf hier de conclusie/analyse m.b.t. de performance eisen en/of gebruiksverwachting t.o.v. de testresultaten. Eventueel in vergelijking met voorgaande testresultaten>").font.size = Pt(
        10)

    #########################
    ### Section: Risico's ###
    #########################
    document.add_heading("Risico\'s", level=1)
    par_risc = document.add_paragraph(style='List Bullet')
    par_risc.add_run("<Vastgestelde risico’s bij in in-productie-name>").font.size = Pt(10)

    ##############################
    ### Section: Vervolgacties ###
    ##############################
    document.add_heading("Vervolgacties", level=1)
    par_vervolg = document.add_paragraph(style='List Bullet')
    par_vervolg.add_run("<De uit te voeren actie bij negatief advies>").font.size = Pt(10)

    #Plaatjes erin gooien

    document.add_page_break()
    document.add_heading('Bijlage A - Plaatjes', level=1)

    #Loop over all images in the image list and add them, with a title

    print(images_list)
    pagecounter = 0

    for image_ in images_list:
        image_split_ = image_.split('/')
        image_split_ = image_split_[2][:-4]

        if os.path.isfile(image_):
            par_images = document.add_paragraph(style='Normal')
            tbl_heading = document.add_heading(str(image_split_), level=2)
            document.add_picture(image_, width=Cm(18.0))
            pagecounter += 1
            if (pagecounter % int(pagecount_number)) == 0:
                document.add_page_break()



    filename = filename.replace('.docx', "_" + str(data) + '_' + datetime.datetime.today().strftime('%d%m%YT%H%M%S') + '.docx')

    working_path = results_dir + '/' + 'report_' + str(data).lower()
    os.makedirs(working_path, exist_ok=True)
    os.chdir(working_path)
    document.save(filename)

    docx_filename_fullpath = working_path + '/' + filename
    os.chdir(home_path)

    return docx_filename_fullpath
